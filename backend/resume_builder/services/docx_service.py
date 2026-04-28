"""
python-docx DOCX generation for resume export.
All docx imports are lazy (inside render_docx) to avoid lxml DLL load errors
at module import time on systems with Application Control policies.
"""
import io
from typing import Any, Dict


def _hex_to_rgb_tuple(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)


def _safe(d: Any, *keys, default="") -> str:
    for k in keys:
        if not isinstance(d, dict):
            return default
        d = d.get(k, default)
    return str(d) if d else default


def render_docx(doc_data: Dict[str, Any]) -> bytes:
    """Render a resume as DOCX bytes. Imports python-docx lazily."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError(
            "python-docx / lxml is not available on this system. "
            "DOCX export is unavailable. Use PDF export instead. "
            f"Details: {exc}"
        )

    content = doc_data.get("content", {})
    personal = content.get("personal", {})
    r, g, b = _hex_to_rgb_tuple(doc_data.get("theme_color", "#6366f1"))
    theme_color = RGBColor(r, g, b)
    size_key = doc_data.get("font_size", "medium")
    base_sizes = {"small": 9, "medium": 10, "large": 12}
    base_size = base_sizes.get(size_key, 10)

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    def add_hr(para):
        p = para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(base_size + 1)
        run.font.color.rgb = theme_color
        add_hr(p)

    # ── Header ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(_safe(personal, "name"))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = theme_color

    contacts = " | ".join(filter(None, [
        _safe(personal, "email"), _safe(personal, "phone"),
        _safe(personal, "location"), _safe(personal, "linkedin"),
    ]))
    cp = doc.add_paragraph(contacts)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
        run.font.size = Pt(9)

    doc.add_paragraph()

    if content.get("summary"):
        section_heading("SUMMARY")
        p = doc.add_paragraph(content["summary"])
        if p.runs: p.runs[0].font.size = Pt(base_size)

    if content.get("experience"):
        section_heading("EXPERIENCE")
        for exp in content["experience"]:
            end = "Present" if exp.get("current") else _safe(exp, "end_date")
            p = doc.add_paragraph()
            r1 = p.add_run(f"{_safe(exp,'title')}  —  {_safe(exp,'company')}")
            r1.bold = True; r1.font.size = Pt(base_size)
            p2 = doc.add_paragraph(f"{_safe(exp,'location')}   {_safe(exp,'start_date')} – {end}")
            if p2.runs: p2.runs[0].font.size = Pt(9)
            for b in exp.get("bullets", []):
                if b.strip():
                    bp = doc.add_paragraph(style="List Bullet")
                    br = bp.add_run(b); br.font.size = Pt(base_size)

    if content.get("education"):
        section_heading("EDUCATION")
        for edu in content["education"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{_safe(edu,'degree')} in {_safe(edu,'field')}")
            r.bold = True; r.font.size = Pt(base_size)
            p2 = doc.add_paragraph(f"{_safe(edu,'institution')}   {_safe(edu,'start_date')} – {_safe(edu,'end_date')}")
            if p2.runs: p2.runs[0].font.size = Pt(9)

    if content.get("skills"):
        section_heading("SKILLS")
        p = doc.add_paragraph(" • ".join(content["skills"]))
        if p.runs: p.runs[0].font.size = Pt(base_size)

    if content.get("projects"):
        section_heading("PROJECTS")
        for proj in content["projects"]:
            p = doc.add_paragraph()
            r = p.add_run(_safe(proj, "name")); r.bold = True; r.font.size = Pt(base_size)
            if _safe(proj, "technologies"):
                p.add_run(f"  [{_safe(proj,'technologies')}]").font.size = Pt(9)
            if _safe(proj, "description"):
                p2 = doc.add_paragraph(_safe(proj, "description"))
                if p2.runs: p2.runs[0].font.size = Pt(base_size)
            for b in proj.get("bullets", []):
                if b.strip():
                    bp = doc.add_paragraph(style="List Bullet")
                    br = bp.add_run(b); br.font.size = Pt(base_size)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
